import os
import json
import re
from datetime import datetime

# ReportLab Imports
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER

# Python-docx Imports
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ================= HELPER FUNCTIONS =================

def convert_md_to_html(text):
    """Converts simple markdown bold, italic, and bullet points to HTML-like tags for ReportLab Paragraphs."""
    # Bold **text** -> <b>text</b>
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Italic *text* -> <i>text</i>
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    return text

def parse_markdown_to_reportlab_story(text, story, styles, primary_color_hex="#1A2E40"):
    """
    Parses a markdown string and appends formatted ReportLab flowables to the story.
    Supports headings (###, ##, #), bullet lists, and paragraphs.
    """
    lines = text.split("\n")
    in_list = False
    
    # Custom heading styles to avoid collisions
    heading_style = ParagraphStyle(
        'PropHeading',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor(primary_color_hex),
        spaceBefore=12,
        spaceAfter=6
    )
    
    body_style = ParagraphStyle(
        'PropBody',
        parent=styles['BodyText'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#2D2D2D'),
        spaceBefore=4,
        spaceAfter=6
    )
    
    bullet_style = ParagraphStyle(
        'PropBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#2D2D2D'),
        leftIndent=15,
        firstLineIndent=-10,
        spaceBefore=2,
        spaceAfter=2
    )

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_list:
                story.append(Spacer(1, 4))
                in_list = False
            continue
            
        # Headings
        if stripped.startswith("### "):
            header_text = stripped[4:]
            story.append(Paragraph(convert_md_to_html(header_text), heading_style))
            in_list = False
        elif stripped.startswith("## "):
            header_text = stripped[3:]
            # Make h2 slightly larger
            h2_style = ParagraphStyle('H2Style', parent=heading_style, fontSize=14, leading=18, spaceBefore=14)
            story.append(Paragraph(convert_md_to_html(header_text), h2_style))
            in_list = False
        elif stripped.startswith("# "):
            header_text = stripped[2:]
            h1_style = ParagraphStyle('H1Style', parent=heading_style, fontSize=16, leading=20, spaceBefore=16)
            story.append(Paragraph(convert_md_to_html(header_text), h1_style))
            in_list = False
            
        # Bullet list item
        elif stripped.startswith("- ") or stripped.startswith("* "):
            bullet_content = stripped[2:]
            story.append(Paragraph(f"&bull; {convert_md_to_html(bullet_content)}", bullet_style))
            in_list = True
            
        # Regular paragraph
        else:
            story.append(Paragraph(convert_md_to_html(stripped), body_style))
            in_list = False

# ================= COMPILER ENGINES =================

def compile_proposal_pdf(output_path, client_name, project_title, proposal_text, freelancer_name, freelancer_email, freelancer_bank):
    """Generates a professional proposal PDF using ReportLab with a deep navy palette."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=54, leftMargin=54,
        topMargin=54, bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # 1. Prominent Header Block / Letterhead (Table layout)
    header_data = [
        [
            Paragraph(f"<b>{freelancer_name}</b><br/><font color='#555555'>{freelancer_email}</font>", ParagraphStyle('HLeft', parent=styles['Normal'], fontSize=11, leading=14)),
            Paragraph("<font size=16 color='#1A2E40'><b>PROJECT PROPOSAL</b></font>", ParagraphStyle('HRight', parent=styles['Normal'], alignment=TA_RIGHT))
        ]
    ]
    header_table = Table(header_data, colWidths=[250, 250])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(header_table)
    
    # Horizontal Divider
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#1A2E40'), spaceBefore=2, spaceAfter=15))
    
    # 2. Metadata / Client Information
    meta_style_left = ParagraphStyle('MLeft', parent=styles['Normal'], fontSize=9, leading=12)
    meta_style_right = ParagraphStyle('MRight', parent=styles['Normal'], fontSize=9, leading=12, alignment=TA_RIGHT)
    
    today_str = datetime.now().strftime("%B %d, %Y")
    meta_data = [
        [
            Paragraph(f"<b>PREPARED FOR:</b><br/>{client_name}", meta_style_left),
            Paragraph(f"<b>DATE:</b> {today_str}<br/><b>PROJECT:</b> {project_title}", meta_style_right)
        ]
    ]
    meta_table = Table(meta_data, colWidths=[250, 250])
    meta_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
    ]))
    story.append(meta_table)
    
    # Divider
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#D0D0D0'), spaceBefore=0, spaceAfter=15))
    
    # 3. Main Proposal Content
    parse_markdown_to_reportlab_story(proposal_text, story, styles, primary_color_hex="#1A2E40")
    
    # 4. Closing Bank Reference (Footer element)
    story.append(Spacer(1, 20))
    footer_data = [
        [
            Paragraph(f"<b>Payment terms & bank details (for reference):</b><br/>{freelancer_bank}", ParagraphStyle('FText', parent=styles['Normal'], fontSize=8, leading=11, textColor=colors.HexColor('#555555')))
        ]
    ]
    footer_table = Table(footer_data, colWidths=[500])
    footer_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F4F6F8')),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING', (0, 0), (-1, -1), 10),
        ('LINEBELOW', (0, 0), (-1, -1), 1, colors.HexColor('#1A2E40')),
    ]))
    story.append(KeepTogether(footer_table))
    
    # Build Document
    doc.build(story)


def compile_proposal_docx(output_path, client_name, project_title, proposal_text, freelancer_name, freelancer_email, freelancer_bank):
    """Generates a professional proposal Word Document (.docx) with structured styling."""
    doc = docx.Document()
    
    # Set Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Letterhead style
    header = doc.add_paragraph()
    h_run1 = header.add_run(f"{freelancer_name}  |  {freelancer_email}\n")
    h_run1.font.name = 'Arial'
    h_run1.font.size = Pt(10)
    h_run1.font.color.rgb = RGBColor(100, 100, 100)
    
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run(f"Project Proposal: {project_title}")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 46, 64) # Navy Accent
    
    # Metadata
    p_meta = doc.add_paragraph()
    p_meta.add_run(f"Prepared For: {client_name}\nDate: {datetime.now().strftime('%B %d, %Y')}\n")
    p_meta.runs[0].font.name = 'Arial'
    p_meta.runs[0].font.size = Pt(11)
    p_meta.runs[0].font.italic = True
    p_meta.runs[0].font.color.rgb = RGBColor(120, 120, 120)
    
    # Divider
    p_div = doc.add_paragraph()
    p_div.add_run("__________________________________________________________________")
    p_div.runs[0].font.color.rgb = RGBColor(200, 200, 200)
    
    # Main content parsing
    lines = proposal_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        if stripped.startswith("### "):
            # Heading 3
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[4:])
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 46, 64)
        elif stripped.startswith("## "):
            # Heading 2
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[3:])
            run.font.name = 'Arial'
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 46, 64)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            # Bullet point
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            # Remove simple markdown bold tags
            content = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped[2:])
            run = p.add_run(content)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
        else:
            # Body text
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
            # Simple bold formatting parser
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = 'Arial'
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(45, 45, 45)
                
    # Footer Bank Details
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(24)
    run_foot = p_foot.add_run(f"Payment instructions: {freelancer_bank}")
    run_foot.font.name = 'Arial'
    run_foot.font.size = Pt(9)
    run_foot.font.italic = True
    run_foot.font.color.rgb = RGBColor(120, 120, 120)
    
    doc.save(output_path)


def compile_invoice_pdf(output_path, invoice_number, created_date, due_date, client_info, project_name, work_items, subtotal, tax_percentage, tax_amount, grand_total, freelancer_payment_details, freelancer_name, freelancer_email):
    """Generates a professional business invoice PDF using ReportLab with a burgundy palette, alternating row colors, and strict bottom payment footer."""
    doc = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        rightMargin=54, leftMargin=54,
        topMargin=54, bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Primary Burgundy hex
    burgundy_hex = "#7A1C1C"
    
    # 1. Header (Letterhead & Title)
    header_data = [
        [
            Paragraph(f"<b>{freelancer_name}</b><br/>{freelancer_email}", ParagraphStyle('HLeftInv', parent=styles['Normal'], fontSize=11, leading=14)),
            Paragraph("<font size=20 color='#7A1C1C'><b>INVOICE</b></font>", ParagraphStyle('HRightInv', parent=styles['Normal'], alignment=TA_RIGHT))
        ]
    ]
    header_table = Table(header_data, colWidths=[250, 250])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(header_table)
    
    # Horizontal Divider (Burgundy)
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor(burgundy_hex), spaceBefore=2, spaceAfter=15))
    
    # 2. Metadata Columns (Client vs Invoice Details)
    meta_style_left = ParagraphStyle('MLeftInv', parent=styles['Normal'], fontSize=9, leading=13)
    meta_style_right = ParagraphStyle('MRightInv', parent=styles['Normal'], fontSize=9, leading=13, alignment=TA_RIGHT)
    
    client_str = f"<b>{client_info.get('client_name', '')}</b><br/>"
    if client_info.get('company_name'):
        client_str += f"{client_info['company_name']}<br/>"
    client_str += f"{client_info.get('contact_details', '')}<br/>{client_info.get('target_email', '')}"
    
    invoice_details_str = f"<b>Invoice Number:</b> {invoice_number}<br/>" \
                          f"<b>Date of Issue:</b> {created_date}<br/>" \
                          f"<b>Due Date:</b> <b>{due_date}</b><br/>" \
                          f"<b>Project:</b> {project_name}"
                          
    meta_data = [
        [
            Paragraph(f"<b>BILLED TO:</b><br/>{client_str}", meta_style_left),
            Paragraph(f"<b>DETAILS:</b><br/>{invoice_details_str}", meta_style_right)
        ]
    ]
    meta_table = Table(meta_data, colWidths=[250, 250])
    meta_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
    ]))
    story.append(meta_table)
    
    # 3. Main Itemized Table (Alternating Row Colors)
    # Headers
    table_data = [
        [
            Paragraph("<b>Work Description</b>", ParagraphStyle('TH1', parent=styles['Normal'], textColor=colors.white, fontName='Helvetica-Bold')),
            Paragraph("<b>Hours</b>", ParagraphStyle('TH2', parent=styles['Normal'], textColor=colors.white, fontName='Helvetica-Bold', alignment=TA_RIGHT)),
            Paragraph("<b>Rate</b>", ParagraphStyle('TH3', parent=styles['Normal'], textColor=colors.white, fontName='Helvetica-Bold', alignment=TA_RIGHT)),
            Paragraph("<b>Line Total</b>", ParagraphStyle('TH4', parent=styles['Normal'], textColor=colors.white, fontName='Helvetica-Bold', alignment=TA_RIGHT))
        ]
    ]
    
    # Append Items
    cell_desc_style = ParagraphStyle('CDesc', parent=styles['Normal'], fontSize=9, leading=12)
    cell_num_style = ParagraphStyle('CNum', parent=styles['Normal'], fontSize=9, leading=12, alignment=TA_RIGHT)
    
    for item in work_items:
        desc = item.get('description', '')
        hours = item.get('hours', 0.0)
        rate = item.get('rate', 0.0)
        line_total = hours * rate
        
        table_data.append([
            Paragraph(desc, cell_desc_style),
            Paragraph(f"{hours:.2f}" if hours > 0 else "Flat", cell_num_style),
            Paragraph(f"${rate:.2f}", cell_num_style),
            Paragraph(f"${line_total:.2f}", cell_num_style)
        ])
        
    # Table Widths: Total 500 pt
    work_table = Table(table_data, colWidths=[270, 70, 75, 85])
    
    # Styling Table with burgundy accents and alternating colors
    t_style = [
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor(burgundy_hex)),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('TOPPADDING', (0, 0), (-1, 0), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E0D0D0')),
    ]
    
    # Alternating row background colors (matching deep burgundy soft tints)
    for i in range(1, len(table_data)):
        bg_color = colors.HexColor('#FDF8F8') if i % 2 == 1 else colors.HexColor('#F5EBEB')
        t_style.append(('BACKGROUND', (0, i), (-1, i), bg_color))
        t_style.append(('BOTTOMPADDING', (0, i), (-1, i), 8))
        t_style.append(('TOPPADDING', (0, i), (-1, i), 8))
        
    work_table.setStyle(TableStyle(t_style))
    story.append(work_table)
    story.append(Spacer(1, 15))
    
    # 4. Summary Block (Subtotal, Tax, Grand Total) - Right Aligned
    summary_style_label = ParagraphStyle('SumLabel', parent=styles['Normal'], fontSize=9.5, leading=13, alignment=TA_RIGHT)
    summary_style_val = ParagraphStyle('SumVal', parent=styles['Normal'], fontSize=9.5, leading=13, alignment=TA_RIGHT)
    summary_style_grand_label = ParagraphStyle('SumGLabel', parent=styles['Normal'], fontSize=11, leading=15, fontName='Helvetica-Bold', textColor=colors.HexColor(burgundy_hex), alignment=TA_RIGHT)
    summary_style_grand_val = ParagraphStyle('SumGVal', parent=styles['Normal'], fontSize=11, leading=15, fontName='Helvetica-Bold', textColor=colors.HexColor(burgundy_hex), alignment=TA_RIGHT)
    
    summary_data = [
        [Paragraph("Subtotal:", summary_style_label), Paragraph(f"${subtotal:.2f}", summary_style_val)],
        [Paragraph(f"Tax ({tax_percentage}%):", summary_style_label), Paragraph(f"${tax_amount:.2f}", summary_style_val)],
        [Paragraph("<b>Grand Total Due:</b>", summary_style_grand_label), Paragraph(f"<b>${grand_total:.2f}</b>", summary_style_grand_val)]
    ]
    
    # Width: 200 pt. Pushed to the right using a leading empty column in a wider table, or a smaller table right-aligned.
    # To place it neatly on the right, we wrap it in a parent table where column 1 is empty spacer.
    outer_summary_data = [
        ["", Table(summary_data, colWidths=[110, 90])]
    ]
    outer_summary_table = Table(outer_summary_data, colWidths=[300, 200])
    outer_summary_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ('TOPPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
    ]))
    story.append(outer_summary_table)
    
    # 5. Payment execution instructions (Lower Margin Box)
    story.append(Spacer(1, 30))
    payment_box_data = [
        [
            Paragraph(f"<font color='{burgundy_hex}'><b>PAYMENT EXECUTION INSTRUCTIONS:</b></font><br/>"
                      f"Please wire or transfer all outstanding dues to:<br/>"
                      f"<b>{freelancer_payment_details}</b><br/>"
                      f"<i>Net due date: {due_date}. Thank you for your business!</i>", 
                      ParagraphStyle('PayText', parent=styles['Normal'], fontSize=8.5, leading=12))
        ]
    ]
    payment_box_table = Table(payment_box_data, colWidths=[500])
    payment_box_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#FDF8F8')),
        ('TOPPADDING', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
        ('LEFTPADDING', (0, 0), (-1, -1), 12),
        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
        ('LINELEFT', (0, 0), (-1, -1), 3, colors.HexColor(burgundy_hex)),
        ('BOX', (0, 0), (-1, -1), 0.5, colors.HexColor('#F0D0D0')),
    ]))
    story.append(KeepTogether(payment_box_table))
    
    # Build Document
    doc.build(story)


def compile_invoice_docx(output_path, invoice_number, created_date, due_date, client_info, project_name, work_items, subtotal, tax_percentage, tax_amount, grand_total, freelancer_payment_details, freelancer_name, freelancer_email):
    """Generates a professional business invoice Word Document (.docx) file."""
    doc = docx.Document()
    
    # Margins
    for s in doc.sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)
        
    # Letterhead header
    p_header = doc.add_paragraph()
    r_hdr = p_header.add_run(f"{freelancer_name}  |  {freelancer_email}")
    r_hdr.font.name = 'Arial'
    r_hdr.font.size = Pt(9.5)
    r_hdr.font.color.rgb = RGBColor(120, 120, 120)
    
    # Title
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(f"Invoice {invoice_number}")
    r_title.font.name = 'Arial'
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(122, 28, 28) # Burgundy accent
    
    # Info Columns
    p_info = doc.add_paragraph()
    info_run = p_info.add_run(
        f"Billed To:\n"
        f"{client_info.get('client_name', '')}\n"
        f"{client_info.get('company_name', '')}\n"
        f"{client_info.get('target_email', '')}\n\n"
        f"Details:\n"
        f"Invoice Date: {created_date}\n"
        f"Due Date: {due_date}\n"
        f"Project: {project_name}"
    )
    info_run.font.name = 'Arial'
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(60, 60, 60)
    
    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1' # Built-in clear table format
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Work Description'
    hdr_cells[1].text = 'Hours'
    hdr_cells[2].text = 'Rate'
    hdr_cells[3].text = 'Line Total'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Arial'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        
    for item in work_items:
        row_cells = table.add_row().cells
        row_cells[0].text = item.get('description', '')
        hours = item.get('hours', 0.0)
        rate = item.get('rate', 0.0)
        line_total = hours * rate
        row_cells[1].text = f"{hours:.2f}" if hours > 0 else "Flat"
        row_cells[2].text = f"${rate:.2f}"
        row_cells[3].text = f"${line_total:.2f}"
        
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.name = 'Arial'
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            
    # Spacing
    doc.add_paragraph()
    
    # Totals block
    p_tot = doc.add_paragraph()
    p_tot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tot_run = p_tot.add_run(
        f"Subtotal: ${subtotal:.2f}\n"
        f"Tax ({tax_percentage}%): ${tax_amount:.2f}\n"
        f"Grand Total: ${grand_total:.2f}"
    )
    tot_run.font.name = 'Arial'
    tot_run.font.size = Pt(11)
    tot_run.font.bold = True
    tot_run.font.color.rgb = RGBColor(122, 28, 28)
    
    # Payment details
    p_pay = doc.add_paragraph()
    p_pay.paragraph_format.space_before = Pt(20)
    pay_run = p_pay.add_run(
        f"Payment Instructions:\n"
        f"Please wire funds to: {freelancer_payment_details}\n"
        f"Payment is due by {due_date}. Thank you!"
    )
    pay_run.font.name = 'Arial'
    pay_run.font.size = Pt(9)
    pay_run.font.italic = True
    pay_run.font.color.rgb = RGBColor(120, 120, 120)
    
    doc.save(output_path)


if __name__ == "__main__":
    # Test suite to verify compile operations
    print("Testing document compiler handlers...")
    
    test_desc = "### Introduction\nThis is a test proposal.\n\n### Deliverables\n- Codebase\n- Tests"
    test_bank = "Test Bank, Acc 9999"
    compile_proposal_pdf("test_proposal.pdf", "Acme Test", "Test Title", test_desc, "Freelancer Jim", "jim@test.com", test_bank)
    compile_proposal_docx("test_proposal.docx", "Acme Test", "Test Title", test_desc, "Freelancer Jim", "jim@test.com", test_bank)
    
    test_items = [{"description": "Item 1", "hours": 5.0, "rate": 100.0}]
    compile_invoice_pdf("test_invoice.pdf", "INV-999", "2026-06-27", "2026-07-04", 
                        {"client_name": "Acme Test", "target_email": "test@acme.com", "contact_details": "123 St"},
                        "Test Project", test_items, 500.0, 10.0, 50.0, 550.0, test_bank, "Freelancer Jim", "jim@test.com")
    compile_invoice_docx("test_invoice.docx", "INV-999", "2026-06-27", "2026-07-04", 
                        {"client_name": "Acme Test", "target_email": "test@acme.com", "contact_details": "123 St"},
                        "Test Project", test_items, 500.0, 10.0, 50.0, 550.0, test_bank, "Freelancer Jim", "jim@test.com")
                        
    print("Test artifacts compiled successfully. Cleaning up...")
    for f in ["test_proposal.pdf", "test_proposal.docx", "test_invoice.pdf", "test_invoice.docx"]:
        if os.path.exists(f):
            os.remove(f)
    print("Cleanup completed.")
